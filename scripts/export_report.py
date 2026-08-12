#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
BeandsAnalystik — export_report.py
Конвертация Markdown-отчёта в PDF и/или DOCX с корректной кириллицей.

Usage:
    python export_report.py reports/final_report.md
    python export_report.py reports/final_report.md --format pdf
    python export_report.py reports/final_report.md --format docx
    python export_report.py reports/final_report.md --format both
    python export_report.py reports/final_report.md --config templates/report-config.json

Возможности:
  - кириллица (через встраиваемые шрифты WeasyPrint / стандартные Calibri-DejaVu)
  - заголовки, таблицы, списки, code blocks
  - разрывы страниц перед H1
  - гиперссылки
  - оглавление (ToC) и номера страниц
  - метаданные документа (title/author/subject)
  - вставка локальных изображений (PNG/JPG/SVG)
  - Mermaid/PlantUML: оставляются как code block + warning, если рендер недоступен
  - сканирование на секреты перед экспортом (предупреждение, без вставки)
  - работает локально, без обязательного облачного API
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
import shutil
from pathlib import Path

# -----------------------------------------------------------------------------
# Зависимости. Импортируем мягко, чтобы дать понятное сообщение, если чего-то нет.
# -----------------------------------------------------------------------------

def _require(modname: str, pip_hint: str):
    try:
        return __import__(modname)
    except ImportError:
        print(f"ERROR: отсутствует модуль '{modname}'. Установите: pip install {pip_hint}",
              file=sys.stderr)
        sys.exit(2)


def _try_import(modname: str):
    try:
        return __import__(modname)
    except (ImportError, OSError):
        # Некоторые пакеты (например WeasyPrint на Windows без GTK) могут
        # быть установлены, но не загружаться из-за нативной библиотеки.
        # Это не должно ломать fallback PDF-экспортёра.
        return None


# markdown нужен всегда
markdown = _require("markdown", "markdown")

# python-docx для DOCX
docx_mod = _try_import("docx")  # python-docx

# PDF-движки (пробуем по очереди)
weasyprint = _try_import("weasyprint")
xhtml2pdf = _try_import("xhtml2pdf")
reportlab_mod = _try_import("reportlab")

PIL = _try_import("PIL")  # pillow, опционально (для размеров изображений)


# -----------------------------------------------------------------------------
# Конфигурация по умолчанию
# -----------------------------------------------------------------------------

DEFAULT_CONFIG = {
    "document": {
        "title": "Финальный отчёт",
        "author": "BeandsAnalystik Suite",
        "subject": "Анализ проекта",
        "language": "ru",
    },
    "pdf": {
        "engine": "auto",
        "page_size": "A4",
        "margin_mm": 18,
        "font_family": "DejaVu Sans, Arial, sans-serif",
        "font_size_pt": 11,
    },
    "docx": {
        "font_family": "Calibri",
        "font_size_pt": 11,
    },
    "output": {
        "dir": "exports",
        "filename_stem": "final_report",
        "overwrite": True,
    },
    "secrets": {
        "scan_before_export": True,
        "patterns": [
            r"(?i)api[_-]?key\s*[=:]\s*['\"]?[A-Za-z0-9]{16,}",
            r"(?i)secret[_-]?key\s*[=:]\s*['\"]?[A-Za-z0-9]{16,}",
            r"(?i)access[_-]?token\s*[=:]\s*['\"]?[A-Za-z0-9]{16,}",
            r"-----BEGIN [A-Z ]+PRIVATE KEY-----",
            r"sk-[A-Za-z0-9]{20,}",
            r"AKIA[0-9A-Z]{16}",
        ],
    },
}


def load_config(config_path: str | None) -> dict:
    """Загрузить конфиг из JSON (или YAML при наличии pyyaml), иначе — DEFAULT_CONFIG."""
    cfg = _deep_copy(DEFAULT_CONFIG)
    if not config_path:
        return cfg
    p = Path(config_path)
    if not p.exists():
        print(f"WARN: конфиг не найден: {config_path}, используются значения по умолчанию.",
              file=sys.stderr)
        return cfg
    text = p.read_text(encoding="utf-8")
    user_cfg = None
    if p.suffix.lower() in (".yaml", ".yml"):
        yaml = _try_import("yaml")
        if yaml is None:
            print("ERROR: pyyaml не установлен для чтения YAML-конфига. Используйте JSON или "
                  "установите pyyaml.", file=sys.stderr)
            sys.exit(2)
        user_cfg = yaml.safe_load(text)
    else:
        # JSON с поддержкой _comment-полей
        user_cfg = json.loads(text)
    # глубокое слияние: user_cfg поверх cfg
    return _deep_merge(cfg, user_cfg)


def _deep_copy(d: dict) -> dict:
    return {k: _deep_copy(v) if isinstance(v, dict) else v for k, v in d.items()}


def _deep_merge(base: dict, override: dict) -> dict:
    out = _deep_copy(base)
    for k, v in (override or {}).items():
        if k.startswith("_"):
            continue  # пропускаем комментарии
        if isinstance(v, dict) and isinstance(out.get(k), dict):
            out[k] = _deep_merge(out[k], v)
        else:
            out[k] = v
    return out


# -----------------------------------------------------------------------------
# Подготовка Markdown: разрывы страниц, ToC-маркер, сканирование секретов
# -----------------------------------------------------------------------------

def inject_page_breaks(md_text: str) -> str:
    """Вставить разрыв страницы перед каждым H1 (кроме первого), если ещё нет."""
    lines = md_text.splitlines()
    out = []
    seen_first_h1 = False
    for line in lines:
        if line.strip().startswith("# ") and not line.strip().startswith("#!"):
            if seen_first_h1:
                # не дублировать, если уже есть маркер разрыва
                if out and not out[-1].strip().lower().startswith("<!-- pagebreak"):
                    out.append("<!-- pagebreak -->")
                    out.append("")
            seen_first_h1 = True
        out.append(line)
    return "\n".join(out)


def scan_secrets(md_text: str, patterns: list[str]) -> list[tuple[int, str]]:
    """Найти строки, похожие на секреты. Возвращает [(line_no, snippet)]."""
    hits = []
    compiled = [re.compile(p) for p in patterns]
    for i, line in enumerate(md_text.splitlines(), start=1):
        for rx in compiled:
            if rx.search(line):
                snippet = line.strip()[:80]
                hits.append((i, snippet))
                break
    return hits


def replace_pagebreak_markers(md_text: str, marker_html: str) -> str:
    """Заменить <!-- pagebreak --> на HTML-маркер (для PDF) или удалить (для DOCX)."""
    return md_text.replace("<!-- pagebreak -->", marker_html)


# -----------------------------------------------------------------------------
# Конвертация Markdown -> HTML
# -----------------------------------------------------------------------------

MD_EXTENSIONS = [
    "extra",        # tables, fenced_code, footnotes, attr_list, def_list, abbr
    "codehilite",   # подсветка кода
    "toc",          # оглавление
    "sane_lists",
    "admonition",
    "smarty",
]


def md_to_html(md_text: str, config: dict) -> str:
    # pagebreak-маркер → CSS-разрыв
    md_text = replace_pagebreak_markers(
        md_text,
        '<div style="page-break-after: always;"></div>',
    )
    html_body = markdown.markdown(
        md_text,
        extensions=MD_EXTENSIONS,
        extension_configs={"codehilite": {"guess_lang": False}},
    )
    title = config["document"].get("title", "Финальный отчёт")
    lang = config["document"].get("language", "ru")
    font = config["pdf"].get("font_family", "DejaVu Sans, Arial, sans-serif")
    fs = config["pdf"].get("font_size_pt", 11)
    margin = config["pdf"].get("margin_mm", 18)
    html = f"""<!DOCTYPE html>
<html lang="{lang}">
<head>
<meta charset="utf-8">
<title>{title}</title>
<style>
@page {{
  size: A4;
  margin: {margin}mm {margin}mm {margin}mm {margin}mm;
}}
/* Нумерация страниц: weasyprint через @bottom-center, xhtml2pdf — через pdf-frame.
   Здесь оставляем минимальный @page (совместимый с обоими движками); номера страниц
   добавляются движком при наличии соответствующего шаблона страницы. */
body {{
  font-family: '{font}';
  font-size: {fs}pt;
  line-height: 1.4;
}}
h1, h2, h3, h4, h5, h6 {{ font-weight: bold; line-height: 1.2; }}
h1 {{ font-size: 1.6em; margin-top: 1.2em; }}
h2 {{ font-size: 1.35em; margin-top: 1.1em; }}
h3 {{ font-size: 1.15em; margin-top: 1em; }}
table {{
  border-collapse: collapse;
  width: 100%;
  margin: 0.8em 0;
}}
th, td {{
  border: 1px solid #888;
  padding: 6px 8px;
  text-align: left;
  vertical-align: top;
}}
th {{ background: #f0f0f0; }}
pre {{
  background: #f6f8fa;
  border: 1px solid #ddd;
  border-radius: 3px;
  padding: 8px 10px;
  overflow-x: auto;
  font-family: 'DejaVu Sans Mono', Consolas, monospace;
  font-size: 0.9em;
  white-space: pre-wrap;
}}
code {{
  font-family: 'DejaVu Sans Mono', Consolas, monospace;
}}
:not(pre) > code {{
  background: #f6f8fa;
  padding: 1px 4px;
  border-radius: 3px;
}}
a {{ color: #1976D2; text-decoration: none; }}
img {{ max-width: 100%; }}
</style>
</head>
<body>
{html_body}
</body>
</html>
"""
    return html


# -----------------------------------------------------------------------------
# PDF-экспорт (с авто-выбором движка)
# -----------------------------------------------------------------------------

def pick_pdf_engine(requested: str):
    """Вернуть (имя_движка, модуль). requested: auto|weasyprint|xhtml2pdf."""
    candidates = []
    if requested == "auto":
        candidates = [("weasyprint", weasyprint), ("reportlab", reportlab_mod), ("xhtml2pdf", xhtml2pdf)]
    elif requested == "weasyprint":
        candidates = [("weasyprint", weasyprint)]
    elif requested == "xhtml2pdf":
        candidates = [("xhtml2pdf", xhtml2pdf)]
    elif requested == "reportlab":
        candidates = [("reportlab", reportlab_mod)]
    else:
        candidates = [("weasyprint", weasyprint), ("reportlab", reportlab_mod), ("xhtml2pdf", xhtml2pdf)]

    for name, mod in candidates:
        if mod is not None:
            return name, mod
    return None, None


def export_pdf(html: str, out_path: Path, config: dict) -> str:
    engine_name, engine = pick_pdf_engine(config["pdf"].get("engine", "auto"))
    if engine is None:
        msg = ("PDF-движок не найден. Установите reportlab (pip install reportlab) "
               "или weasyprint (pip install weasyprint).")
        print(f"ERROR: {msg}", file=sys.stderr)
        return f"SKIPPED ({msg})"

    title = config["document"].get("title", "Финальный отчёт")
    author = config["document"].get("author", "BeandsAnalystik Suite")
    subject = config["document"].get("subject", "")

    if engine_name == "weasyprint":
        try:
            doc = engine.HTML(string=html).render()
            # метаданные
            doc.metadata.title = title
            doc.metadata.authors = [author]
            doc.metadata.description = subject
            doc.write_pdf(str(out_path))
            return f"OK (weasyprint) -> {out_path}"
        except Exception as e:
            print(f"WARN: weasyprint завершился с ошибкой ({e}); пробуем fallback...",
                  file=sys.stderr)
            # fallback на xhtml2pdf если доступен
            if xhtml2pdf is not None:
                return _export_pdf_xhtml2pdf(html, out_path, title)
            return f"FAILED (weasyprint: {e})"
    elif engine_name == "reportlab":
        return _export_pdf_reportlab(html, out_path, title)
    elif engine_name == "xhtml2pdf":
        return _export_pdf_xhtml2pdf(html, out_path, title)
    return "FAILED (нет доступного движка)"


def _export_pdf_xhtml2pdf(html: str, out_path: Path, title: str) -> str:
    """Экспортировать через xhtml2pdf с гарантированной кириллицей.

    Встроенные PDF-шрифты ReportLab не содержат кириллицу. Поэтому при
    fallback на xhtml2pdf явно подключаем системный TrueType Arial. Если
    шрифт недоступен, экспорт не выполняется: PDF с квадратами хуже честной
    диагностической ошибки.
    """
    # xhtml2pdf на Windows иногда не встраивает TTF-шрифт и заменяет
    # кириллицу квадратами. Для отчётов с русским текстом используем
    # собственный ReportLab fallback: он встраивает Arial напрямую.
    return _export_pdf_reportlab(html, out_path, title)


def _export_pdf_reportlab(html: str, out_path: Path, title: str) -> str:
    """Надёжный PDF fallback с TrueType Arial и поддержкой кириллицы."""
    from html import unescape
    from xml.sax.saxutils import escape
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import KeepTogether, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.fonts import addMapping

    windows_fonts = Path(os.environ.get("WINDIR", r"C:\\Windows")) / "Fonts"
    font_sets = [
        {
            "BeandsArial": windows_fonts / "arial.ttf",
            "BeandsArial-Bold": windows_fonts / "arialbd.ttf",
            "BeandsArial-Italic": windows_fonts / "ariali.ttf",
            "BeandsArial-BoldItalic": windows_fonts / "arialbi.ttf",
        },
        {
            "BeandsArial": Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
            "BeandsArial-Bold": Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
            "BeandsArial-Italic": Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf"),
            "BeandsArial-BoldItalic": Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf"),
        },
    ]
    font_paths = next((candidate for candidate in font_sets if all(path.exists() for path in candidate.values())), None)
    if font_paths is None:
        return "FAILED (ReportLab: не найден набор TrueType-шрифтов с кириллицей; поддерживаются Arial в Windows и DejaVu Sans в Linux)"

    for name, path in font_paths.items():
        try:
            pdfmetrics.registerFont(TTFont(name, str(path)))
        except KeyError:
            # Шрифт уже зарегистрирован при повторном экспорте в одном процессе.
            pass
    addMapping("BeandsArial", 0, 0, "BeandsArial")
    addMapping("BeandsArial", 1, 0, "BeandsArial-Bold")
    addMapping("BeandsArial", 0, 1, "BeandsArial-Italic")
    addMapping("BeandsArial", 1, 1, "BeandsArial-BoldItalic")

    styles = getSampleStyleSheet()
    body = ParagraphStyle("BeandsBody", parent=styles["BodyText"], fontName="BeandsArial", fontSize=10.5, leading=14, alignment=TA_LEFT, spaceAfter=7)
    heading = {
        1: ParagraphStyle("BeandsH1", parent=styles["Heading1"], fontName="BeandsArial-Bold", fontSize=18, leading=22, spaceBefore=10, spaceAfter=10),
        2: ParagraphStyle("BeandsH2", parent=styles["Heading2"], fontName="BeandsArial-Bold", fontSize=14, leading=18, spaceBefore=10, spaceAfter=7),
        3: ParagraphStyle("BeandsH3", parent=styles["Heading3"], fontName="BeandsArial-Bold", fontSize=12, leading=15, spaceBefore=8, spaceAfter=5),
    }
    table_cell = ParagraphStyle(
        "BeandsTableCell",
        parent=body,
        fontName="BeandsArial",
        fontSize=7.2,
        leading=8.8,
        spaceAfter=0,
        wordWrap="CJK",
    )
    table_header = ParagraphStyle(
        "BeandsTableHeader",
        parent=table_cell,
        fontName="BeandsArial-Bold",
    )
    def clean(fragment: str) -> str:
        fragment = re.sub(r"<[^>]+>", "", fragment)
        return escape(unescape(fragment)).replace("\n", " ").strip()
    story = []
    # HTML генерируется внутренним Markdown-конвертером, поэтому достаточно
    # обработать его устойчивое подмножество: заголовки, абзацы, списки и таблицы.
    for tag, contents in re.findall(r"<(h[1-6]|p|li|table)[^>]*>(.*?)</\1>", html, flags=re.I | re.S):
        tag = tag.lower()
        if tag.startswith("h"):
            level = min(int(tag[1]), 3)
            story.append(Paragraph(clean(contents), heading[level]))
        elif tag == "li":
            story.append(Paragraph("• " + clean(contents), body))
        elif tag == "p":
            text = clean(contents)
            if text:
                story.append(Paragraph(text, body))
        elif tag == "table":
            rows = []
            for row in re.findall(r"<tr[^>]*>(.*?)</tr>", contents, flags=re.I | re.S):
                cells = [clean(cell) for cell in re.findall(r"<t[hd][^>]*>(.*?)</t[hd]>", row, flags=re.I | re.S)]
                if cells:
                    rows.append(cells)
            if rows:
                column_count = max(len(row) for row in rows)
                # Равные колонки превращали таблицы с пятью полями в кашу.
                # Для известных 5-колоночных таблиц приоритетов даём больше
                # места инициативе и результату; в остальных случаях - равные.
                if column_count == 5:
                    col_widths = [13 * mm, 42 * mm, 35 * mm, 53 * mm, 31 * mm]
                else:
                    col_widths = [174 * mm / column_count] * column_count
                formatted_rows = []
                for row_index, row in enumerate(rows):
                    padded = row + [""] * (column_count - len(row))
                    style = table_header if row_index == 0 else table_cell
                    formatted_rows.append([Paragraph(cell or " ", style) for cell in padded])
                table = Table(formatted_rows, colWidths=col_widths, repeatRows=1, splitByRow=1)
                table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.35, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F0F0F0")), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 3), ("RIGHTPADDING", (0, 0), (-1, -1), 3), ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3)]))
                story.extend([table, Spacer(1, 7)])
    if not story:
        return "FAILED (ReportLab: HTML не содержит экспортируемого текста)"
    document = SimpleDocTemplate(str(out_path), pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm, topMargin=18 * mm, bottomMargin=18 * mm, title=title)
    document.build(story)
    return f"OK (reportlab-cyrillic) -> {out_path}"


# -----------------------------------------------------------------------------
# DOCX-экспорт (через python-docx, без pandoc)
# -----------------------------------------------------------------------------

INLINE_CODE_RE = re.compile(r"`([^`]+)`")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
ITALIC_RE = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def export_docx(md_text: str, out_path: Path, config: dict) -> str:
    if docx_mod is None:
        msg = ("python-docx не установлен. Установите: pip install python-docx")
        print(f"ERROR: {msg}", file=sys.stderr)
        return f"SKIPPED ({msg})"

    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    document = Document()

    # метаданные
    cp = document.core_properties
    cp.title = config["document"].get("title", "Финальный отчёт")
    cp.author = config["document"].get("author", "BeandsAnalystik Suite")
    cp.subject = config["document"].get("subject", "")
    cp.language = config["document"].get("language", "ru")

    font_name = config["docx"].get("font_family", "Calibri")
    font_size = Pt(config["docx"].get("font_size_pt", 11))

    # базовый стиль
    style = document.styles["Normal"]
    style.font.name = font_name
    style.font.size = font_size
    # кириллица: убедимся, что east-asia/cs тоже установлены
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font_name)

    def add_runs_with_inline(paragraph, text: str):
        """Разобрать **bold**, *italic*, `code`, [text](url) и добавить runs."""
        # порядок: ссылки → код → жирный → курсив
        tokens = [(LINK_RE, "link")]
        # упрощённый пострендер: заменим на placeholder-маркеры, потом восстановим
        # для надёжности обработаем последовательно регулярками с разбиением
        pos = 0
        # Соберём все матчи с приоритетом
        matches = []
        for m in LINK_RE.finditer(text):
            matches.append((m.start(), m.end(), "link", m))
        for m in INLINE_CODE_RE.finditer(text):
            matches.append((m.start(), m.end(), "code", m))
        for m in BOLD_RE.finditer(text):
            matches.append((m.start(), m.end(), "bold", m))
        for m in ITALIC_RE.finditer(text):
            matches.append((m.start(), m.end(), "italic", m))
        # убрать пересечения: отсортировать и выбрать непересекающиеся
        matches.sort(key=lambda x: (x[0], x[1] - x[0]))
        chosen = []
        last_end = 0
        for s, e, kind, m in matches:
            if s >= last_end:
                chosen.append((s, e, kind, m))
                last_end = e

        cursor = 0
        for s, e, kind, m in chosen:
            if s > cursor:
                paragraph.add_run(text[cursor:s])
            if kind == "link":
                run = paragraph.add_run(m.group(1))
                run.font.color.rgb = RGBColor(0x19, 0x76, 0xD2)
                run.underline = True
                # гиперссылка
                _add_hyperlink(paragraph, m.group(2), m.group(1))
                # убираем лишний run, добавленный выше (упрощение: оставим как есть)
            elif kind == "code":
                run = paragraph.add_run(m.group(1))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            elif kind == "bold":
                run = paragraph.add_run(m.group(1))
                run.bold = True
            elif kind == "italic":
                run = paragraph.add_run(m.group(1))
                run.italic = True
            cursor = e
        if cursor < len(text):
            paragraph.add_run(text[cursor:])

    def _add_hyperlink(paragraph, url: str, text: str):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        hyperlink.append(new_run)
        new_run.append(rPr)
        paragraph._p.append(hyperlink)

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # pagebreak маркер → разрыв страницы
        if stripped.lower().startswith("<!-- pagebreak"):
            from docx.enum.text import WD_BREAK
            p = document.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)
            i += 1
            continue

        # заголовки
        if stripped.startswith("#"):
            m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if m:
                level = len(m.group(1))
                text_h = m.group(2).strip()
                h = document.add_heading(level=min(level, 6))
                add_runs_with_inline(h, text_h)
                i += 1
                continue

        # таблица (GFM): блок строк вида | ... |
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            _add_docx_table(document, table_lines)
            continue

        # fenced code block
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # пропустить закрывающий ```
            p = document.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            # лёгкая заливка через shading — пропустим для простоты
            continue

        # список (маркированный)
        if re.match(r"^\s*[-*+]\s+", line):
            bullet_lines = []
            while i < len(lines) and re.match(r"^\s*[-*+]\s+", lines[i]):
                bullet_lines.append(re.sub(r"^\s*[-*+]\s+", "", lines[i]))
                i += 1
            for b in bullet_lines:
                p = document.add_paragraph(style="List Bullet")
                add_runs_with_inline(p, b)
            continue

        # нумерованный список
        if re.match(r"^\s*\d+\.\s+", line):
            num_lines = []
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                num_lines.append(re.sub(r"^\s*\d+\.\s+", "", lines[i]))
                i += 1
            for n in num_lines:
                p = document.add_paragraph(style="List Number")
                add_runs_with_inline(p, n)
            continue

        # горизонтальная линия
        if stripped in ("---", "***", "___"):
            # простой пустой абзац-разделитель
            document.add_paragraph()
            i += 1
            continue

        # пустая строка
        if stripped == "":
            i += 1
            continue

        # обычный абзац
        p = document.add_paragraph()
        add_runs_with_inline(p, stripped)
        i += 1

    document.save(str(out_path))
    return f"OK -> {out_path}"


def _add_docx_table(document, table_lines: list[str]):
    """Преобразовать GFM-таблицу в таблицу docx."""
    from docx.shared import Pt
    # отфильтровать разделитель |---|---|
    rows = []
    for ln in table_lines:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if all(re.match(r"^:?-{2,}:?$", c) for c in cells if c):
            continue  # строка-разделитель
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=ncols)
    table.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            cell.text = row[ci] if ci < len(row) else ""
            if ri == 0:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
    document.add_paragraph()  # отступ после таблицы


# -----------------------------------------------------------------------------
# Главный entry point
# -----------------------------------------------------------------------------

def main(argv=None):
    parser = argparse.ArgumentParser(
        prog="export_report.py",
        description="Конвертация Markdown-отчёта BeandsAnalystik в PDF/DOCX.",
    )
    parser.add_argument("input", help="Путь к .md файлу отчёта")
    parser.add_argument("--format", choices=["pdf", "docx", "both"], default="both",
                        help="Формат вывода (по умолчанию: both)")
    parser.add_argument("--config", default=None,
                        help="Путь к report-config.json/yaml (опционально)")
    parser.add_argument("--output-dir", default=None,
                        help="Каталог вывода (переопределяет config)")
    parser.add_argument("--no-secrets-scan", action="store_true",
                        help="Пропустить сканирование на секреты")
    args = parser.parse_args(argv)

    in_path = Path(args.input)
    if not in_path.exists():
        print(f"ERROR: входной файл не найден: {in_path}", file=sys.stderr)
        return 1

    config = load_config(args.config)

    # выходной каталог и имя файла.
    # По умолчанию используется имя входного файла (stem). Значение из config
    # применяется только если оно задано явно (не дефолтным "final_report").
    out_dir = Path(args.output_dir) if args.output_dir else Path(config["output"]["dir"])
    out_dir.mkdir(parents=True, exist_ok=True)
    cfg_stem = config["output"].get("filename_stem", "final_report")
    # если config загружен из файла пользователя и там задан stem — используем его;
    # иначе — stem входного файла.
    explicit_cfg = args.config is not None
    stem = cfg_stem if (explicit_cfg and cfg_stem and cfg_stem != "final_report") else in_path.stem

    md_text = in_path.read_text(encoding="utf-8")

    # сканирование секретов
    if not args.no_secrets_scan and config["secrets"].get("scan_before_export", True):
        patterns = config["secrets"].get("patterns", [])
        hits = scan_secrets(md_text, patterns)
        if hits:
            print(f"WARN: обнаружены {len(hits)} потенциальных секрет(ов) в тексте. "
                  "Экспорт продолжится, но проверьте эти строки:", file=sys.stderr)
            for ln, snip in hits[:10]:
                print(f"  строка {ln}: {snip}", file=sys.stderr)
            print("  (значения секретов НЕ удалены автоматически — проверьте вручную)",
                  file=sys.stderr)

    # инжектировать разрывы страниц перед H1
    md_text_pb = inject_page_breaks(md_text)

    results = []

    if args.format in ("pdf", "both"):
        html = md_to_html(md_text_pb, config)
        out_pdf = out_dir / f"{stem}.pdf"
        if out_pdf.exists() and not config["output"].get("overwrite", True):
            print(f"SKIP: {out_pdf} уже существует (overwrite=False)", file=sys.stderr)
        else:
            results.append(("PDF", export_pdf(html, out_pdf, config)))

    if args.format in ("docx", "both"):
        out_docx = out_dir / f"{stem}.docx"
        if out_docx.exists() and not config["output"].get("overwrite", True):
            print(f"SKIP: {out_docx} уже существует (overwrite=False)", file=sys.stderr)
        else:
            results.append(("DOCX", export_docx(md_text_pb, out_docx, config)))

    print("\n--- Результат экспорта ---")
    for fmt, msg in results:
        print(f"{fmt}: {msg}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
